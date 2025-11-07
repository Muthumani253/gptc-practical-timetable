# export_word.py — Subject-wise Word/PDF export (all batches combined)
# Alignment-optimized: fixed column widths, grid style, cell margins, Times New Roman
# NOTE: This version DOES NOT include "COURSE CODE" or "NAME OF THE COURSE" lines.
# Requires: python-docx  (and docx2pdf for PDF on Windows + MS Word)

from io import BytesIO
from datetime import datetime
import os
import tempfile
import pandas as pd

from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml.shared import OxmlElement, qn

# Optional PDF (Windows + MS Word installed)
try:
    from docx2pdf import convert as docx2pdf_convert
    DOCX2PDF_AVAILABLE = True
except Exception:
    DOCX2PDF_AVAILABLE = False

from scheduler_logic import (
    init_db, list_practicals_by, get_batches, list_batch_members,
    get_students_for_practical, list_assigned_reg_nos_for_practical
)

init_db()

# ------------------- CONFIG -------------------
# If you want to maintain a base template docx (with logo, header/footer),
# set TEMPLATE_PATH to that .docx file. If None, a fresh document is created.
# Example: TEMPLATE_PATH = r"F:\GPTC_Practical_Timetable\app\templates\sign_sheet.docx"
TEMPLATE_PATH = None

# ---------- Helpers ----------
def _fmt_ampm(hhmm: str) -> str:
    try:
        return pd.to_datetime(hhmm, format="%H:%M").strftime("%I:%M %p").lstrip("0")
    except Exception:
        return hhmm

def _roman(n: int) -> str:
    ROMANS = ["","I","II","III","IV","V","VI","VII","VIII","IX","X","XI","XII"]
    return ROMANS[n] if 0 <= n < len(ROMANS) else str(n)

def _cell_set_margins(cell, **margins_twips):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    # remove existing margins to avoid duplicates
    for el in tcPr.findall(qn('w:tcMar')):
        tcPr.remove(el)
    tcMar = OxmlElement('w:tcMar')
    for k, v in margins_twips.items():
        node = OxmlElement(f'w:{k}')
        node.set(qn('w:w'), str(v))   # twips
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def _para(doc, text, bold=False, size=12, align="L", after_pt=2):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = bold
    r.font.size = Pt(size)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if align == "C" else WD_ALIGN_PARAGRAPH.LEFT
    if after_pt:
        p.space_after = Pt(after_pt)
    return p

def _apply_cell_text(cell, text, *, align="L", size=10, bold=False, valign="M"):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = {"L": WD_ALIGN_PARAGRAPH.LEFT,
                   "C": WD_ALIGN_PARAGRAPH.CENTER,
                   "R": WD_ALIGN_PARAGRAPH.RIGHT}.get(align, WD_ALIGN_PARAGRAPH.LEFT)
    r = p.add_run(text)
    r.bold = bold
    r.font.size = Pt(size)
    cell.vertical_alignment = {
        "T": WD_CELL_VERTICAL_ALIGNMENT.TOP,
        "M": WD_CELL_VERTICAL_ALIGNMENT.CENTER,
        "B": WD_CELL_VERTICAL_ALIGNMENT.BOTTOM
    }.get(valign, WD_CELL_VERTICAL_ALIGNMENT.CENTER)

# ---------- Core build ----------
def build_subject_docx_bytes(practical_code: str, header_overrides: dict | None = None) -> bytes:
    """
    Create one DOCX for the subject combining all its batches, in the official format.
    header_overrides keys (optional):
      session_title, institute_line, department_line,
      subject_code, subject_name, practical_code, year_sem, date_line
    NOTE: This builder intentionally omits "COURSE CODE" and "NAME OF THE COURSE" lines.
    """
    header_overrides = header_overrides or {}

    pm = list_practicals_by()
    row = pm[pm["practical_code"] == practical_code]
    if row.empty:
        raise ValueError("Invalid practical code.")

    # Defaults from data
    subj_name = row["subject_name"].iloc[0]
    subj_code = row["sub_code"].iloc[0] if "sub_code" in row.columns else ""
    dept_name = row["dept_name"].iloc[0] if "dept_name" in row.columns else ""
    year_sem_default = row["year_sem"].iloc[0] if "year_sem" in row.columns else ""
    institute_line_guess = row["institute_line"].iloc[0] if "institute_line" in row.columns else ""

    # Gather batches & members
    batches = get_batches(practical_code).sort_values(["date", "start_time"]).reset_index(drop=True)

    # Build date line: single date or range
    if batches.empty:
        date_line_auto = ""
    else:
        dates = batches["date"].dropna().tolist()
        dmin, dmax = min(dates), max(dates)
        date_line_auto = dmin if dmin == dmax else f"{dmin} to {dmax}"

    # Header fields
    session_title = header_overrides.get("session_title") or f"BOARD PRACTICAL EXAMINATIONS {datetime.now().strftime('%B').upper()} – {datetime.now().year}"
    institute_line = header_overrides.get("institute_line") or institute_line_guess
    department_line = header_overrides.get("department_line") or f"DEPARTMENT OF {dept_name}".upper()
    subject_code = header_overrides.get("subject_code") or subj_code
    subject_name = header_overrides.get("subject_name") or subj_name
    practical_code_str = header_overrides.get("practical_code") or practical_code
    year_sem = header_overrides.get("year_sem") or year_sem_default
    date_line = header_overrides.get("date_line") or date_line_auto

    # Document (optionally based on TEMPLATE_PATH)
    if TEMPLATE_PATH and os.path.exists(TEMPLATE_PATH):
        doc = Document(TEMPLATE_PATH)
    else:
        doc = Document()
    # A4 & margins
    section = doc.sections[0]
    section.page_height = Cm(29.7)
    section.page_width = Cm(21.0)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)

    # Default font
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(11)

    # Top headings
    _para(doc, session_title, bold=True, size=14, align="C", after_pt=2)
    if institute_line:
        _para(doc, institute_line, bold=True, size=12, align="C", after_pt=1)
    if department_line:
        _para(doc, department_line, bold=True, size=12, align="C", after_pt=10)

    # NOTE: Removed COURSE CODE and NAME OF THE COURSE lines per request

    # Subject details
    def det(label, val): _para(doc, f"{label} : {val}", size=11, bold=False, align="L", after_pt=2)

    det("SUBJECT CODE", subject_code)
    det("NAME OF THE SUBJECT", subject_name)
    det("PRACTICAL CODE", practical_code_str)
    if year_sem:
        det("YEAR / SEMESTER", year_sem)
    if date_line:
        det("DATE", date_line)

    # TIMETABLE header
    _para(doc, "TIMETABLE", bold=True, size=12, align="C", after_pt=6)

    # Table: fixed layout
    table = doc.add_table(rows=1, cols=5)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False

    # Column widths (sum approx inner width)
    col_widths = [Cm(1.2), Cm(1.8), Cm(4.2), Cm(7.2), Cm(1.5)]
    for i, w in enumerate(col_widths):
        table.columns[i].width = w
        table.rows[0].cells[i].width = w

    # Header row
    hdr = table.rows[0].cells
    headers = ["S.NO", "BATCH NO", "DATE & TIME", "REG.NO OF STUDENTS", "TOTAL NO OF STUDENTS"]
    for i, t in enumerate(headers):
        _apply_cell_text(hdr[i], t, align="C", size=10, bold=True, valign="M")
        _cell_set_margins(hdr[i], top="80", start="90", bottom="80", end="90")

    # Body rows
    sn = 1
    grand_total = 0
    for _, b in batches.iterrows():
        batch_id = int(b["batch_id"])
        mem = list_batch_members(batch_id, detailed=True)
        reg_list = mem["reg_no"].tolist() if not mem.empty else []
        reg_text = ", ".join(str(x) for x in reg_list)

        row_cells = table.add_row().cells
        for i, w in enumerate(col_widths):
            row_cells[i].width = w

        _apply_cell_text(row_cells[0], str(sn), align="C", size=10)
        _apply_cell_text(row_cells[1], _roman(int(b["batch_no"])), align="C", size=10)
        dt = f"{b['date']} & {_fmt_ampm(b['start_time'])} – {_fmt_ampm(b['end_time'])}"
        _apply_cell_text(row_cells[2], dt, align="C", size=10)
        _apply_cell_text(row_cells[3], reg_text, align="L", size=10)
        cnt = len(reg_list)
        _apply_cell_text(row_cells[4], f"{cnt:02d}", align="C", size=10)

        for c in row_cells:
            _cell_set_margins(c, top="80", start="90", bottom="80", end="90")

        grand_total += cnt
        sn += 1

    # Total row
    tot_row = table.add_row().cells
    for i, w in enumerate(col_widths):
        tot_row[i].width = w
    _apply_cell_text(tot_row[0], "", align="C", size=10)
    _apply_cell_text(tot_row[1], "", align="C", size=10)
    _apply_cell_text(tot_row[2], "TOTAL", align="C", size=10, bold=True)
    _apply_cell_text(tot_row[3], "", align="L", size=10)
    _apply_cell_text(tot_row[4], f"{grand_total:02d}", align="C", size=10, bold=True)
    for c in tot_row:
        _cell_set_margins(c, top="80", start="90", bottom="80", end="90")

     # Signature row
    # Add 3 blank lines before the signature table
    for _ in range(2):
        _para(doc, "", size=12, after_pt=0)

    _para(doc, "", size=6, after_pt=2)
    sig = doc.add_table(rows=1, cols=3)
    sig.style = 'Table Grid'
    sig.alignment = WD_TABLE_ALIGNMENT.CENTER
    sig.autofit = True
    row_sig = sig.rows[0]

    _apply_cell_text(row_sig.cells[0], "INTERNAL EXAMINER", align="C", size=10, bold=True)
    _apply_cell_text(row_sig.cells[1], "HOD", align="C", size=10, bold=True)
    _apply_cell_text(row_sig.cells[2], "CHIEF SUPERINTENDENT", align="C", size=10, bold=True)

    for cell in row_sig.cells:
        _cell_set_margins(cell, top="120", start="120", bottom="120", end="120")


    # Output bytes
    bio = BytesIO()
    doc.save(bio)
    return bio.getvalue()


def try_convert_docx_to_pdf(docx_bytes: bytes) -> bytes | None:
    """Try to convert DOCX to PDF using docx2pdf (Windows + Word). Returns bytes or None."""
    if not DOCX2PDF_AVAILABLE:
        return None
    with tempfile.TemporaryDirectory() as td:
        docx_path = os.path.join(td, "out.docx")
        pdf_path = os.path.join(td, "out.pdf")
        with open(docx_path, "wb") as f:
            f.write(docx_bytes)
        try:
            docx2pdf_convert(docx_path, pdf_path)
            with open(pdf_path, "rb") as f:
                return f.read()
        except Exception:
            return None
